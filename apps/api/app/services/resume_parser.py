from pathlib import Path
import re

import fitz
from docx import Document


class ResumeExtractionError(Exception):
    """Raised when a supported resume file cannot be read."""


def extract_resume_text(file_path: Path) -> str:
    """Extract normalized text from a PDF or DOCX resume."""
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            with fitz.open(file_path) as document:
                text = "\n".join(page.get_text("text") for page in document)
        elif suffix == ".docx":
            document = Document(file_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
            text = "\n".join(paragraphs + table_cells)
        else:
            raise ResumeExtractionError("This resume format is not supported.")
    except (fitz.FileDataError, ValueError, OSError) as error:
        raise ResumeExtractionError("We couldn't read this resume file.") from error

    normalized_text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not normalized_text:
        raise ResumeExtractionError(
            "No readable text was found. If this is a scanned resume, use a text-based PDF."
        )

    return normalized_text
